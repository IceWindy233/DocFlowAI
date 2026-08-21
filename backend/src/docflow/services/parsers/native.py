from __future__ import annotations

import hashlib
import importlib.util
import shutil
import subprocess
import tempfile
from pathlib import Path

import pdfplumber
from docx import Document as WordDocument
from openpyxl import load_workbook
from PIL import Image
from pypdf import PdfReader

from docflow.domain.documents import BoundingBox, PageV1, TableV1
from docflow.services.parsers.base import ParseContext, ParsedContent, ParserError
from docflow.services.parsers.common import (
    classify_page,
    quality_score,
    render_pdf_page,
    render_quicklook_thumbnail,
    stable_page_id,
    table_from_rows,
)
from docflow.services.parsers.ocr import OcrEngine


def _artifact_page_prefix(context: ParseContext, page_number: int) -> Path:
    root = getattr(context.artifacts, "root", Path(tempfile.gettempdir()))
    return Path(root) / context.document_id / "pages" / f"page-{page_number:04d}"


def _convert_with_libreoffice(source: Path, output_dir: Path, target: str) -> Path:
    executable = shutil.which("soffice")
    if not executable:
        mac_path = Path("/Applications/LibreOffice.app/Contents/MacOS/soffice")
        executable = str(mac_path) if mac_path.exists() else None
    if not executable:
        raise ParserError("未找到 LibreOffice soffice")
    output_dir.mkdir(parents=True, exist_ok=True)
    process = subprocess.run(
        [
            executable,
            "--headless",
            "--convert-to",
            target,
            "--outdir",
            str(output_dir),
            str(source),
        ],
        check=False,
        capture_output=True,
        text=True,
        timeout=300,
    )
    if process.returncode != 0:
        raise ParserError(f"LibreOffice 转换失败：{process.stderr[-300:]}")
    candidates = list(output_dir.glob(f"{source.stem}.*"))
    expected_suffix = f".{target.split(':', 1)[0]}"
    converted = next((item for item in candidates if item.suffix.lower() == expected_suffix), None)
    if converted is None:
        raise ParserError("LibreOffice 未生成预期文件")
    return converted


def _pdf_tables(page: pdfplumber.page.Page, document_id: str, page_number: int) -> list[TableV1]:
    tables: list[TableV1] = []
    try:
        found = page.find_tables()
    except Exception:
        return tables
    for table_number, found_table in enumerate(found, start=1):
        extracted = found_table.extract() or []
        rows = [["" if value is None else str(value).strip() for value in row] for row in extracted]
        if not rows:
            continue
        table = table_from_rows(f"{document_id}_p{page_number:04d}_t{table_number:03d}", rows)
        table.bbox = BoundingBox(
            x0=float(found_table.bbox[0]),
            y0=float(found_table.bbox[1]),
            x1=float(found_table.bbox[2]),
            y1=float(found_table.bbox[3]),
            coordinate_space="pdf_point",
        )
        for cell, bbox in zip(table.cells, found_table.cells, strict=False):
            if bbox:
                cell.bbox = BoundingBox(
                    x0=float(bbox[0]),
                    y0=float(bbox[1]),
                    x1=float(bbox[2]),
                    y1=float(bbox[3]),
                    coordinate_space="pdf_point",
                )
        tables.append(table)
    return tables


def _crop_table_screenshot(
    page_image: Path,
    table: TableV1,
    page_width: float,
    page_height: float,
    output: Path,
) -> Path | None:
    if table.bbox is None or page_width <= 0 or page_height <= 0:
        return None
    try:
        with Image.open(page_image) as image:
            scale_x = image.width / page_width
            scale_y = image.height / page_height
            crop_box = (
                max(0, int(table.bbox.x0 * scale_x)),
                max(0, int(table.bbox.y0 * scale_y)),
                min(image.width, int(table.bbox.x1 * scale_x)),
                min(image.height, int(table.bbox.y1 * scale_y)),
            )
            output.parent.mkdir(parents=True, exist_ok=True)
            image.crop(crop_box).save(output, format="PNG")
        return output
    except (OSError, ValueError):
        return None


class PdfParser:
    name = "pypdf+ocr"
    supported_extensions = {".pdf"}

    def parse(self, context: ParseContext) -> ParsedContent:
        try:
            reader = PdfReader(str(context.source_path), strict=False)
            plumber = pdfplumber.open(context.source_path)
        except Exception as exc:
            raise ParserError(f"PDF 打开失败：{exc}") from exc
        if len(reader.pages) > context.config.parsing.max_page_count:
            raise ParserError("PDF 页数超过配置上限")
        pages: list[PageV1] = []
        warnings: list[str] = []
        ocr = OcrEngine()
        try:
            for index, pdf_page in enumerate(reader.pages, start=1):
                try:
                    native_text = pdf_page.extract_text() or ""
                except Exception:
                    native_text = ""
                plumber_page = plumber.pages[index - 1]
                tables = _pdf_tables(plumber_page, context.document_id, index)
                score, signals = quality_score(native_text)
                needs_ocr = (
                    len(native_text.strip()) < context.config.parsing.searchable_chars_per_page_min
                )
                complex_table = any(table.complex for table in tables)
                image_path: Path | None = None
                parser_route = "NATIVE_PDF"
                text = native_text
                if needs_ocr or complex_table or context.config.indexes.visual_enabled:
                    image_path = render_pdf_page(
                        context.source_path,
                        index,
                        _artifact_page_prefix(context, index),
                        context.config.parsing.pdf_render_dpi,
                    )
                if needs_ocr and image_path:
                    ocr_text, engine = ocr.recognize(image_path)
                    if ocr_text:
                        text = ocr_text
                        parser_route = f"SCANNED_PDF_OCR:{engine}"
                        score, signals = quality_score(text)
                    else:
                        warnings.append(f"第 {index} 页 OCR 不可用或无结果")
                page_type = classify_page(text, tables, image_only=needs_ocr)
                visual_required = (
                    complex_table
                    or page_type in {"MIXED", "STAMPED"}
                    # OCR quality only describes extracted text. A scanned page
                    # can still carry layout, seals, signatures and visual table
                    # structure that text retrieval cannot preserve, so every
                    # scan must enter the ColPali index.
                    or page_type == "SCAN"
                )
                if visual_required and image_path is None:
                    image_path = render_pdf_page(
                        context.source_path,
                        index,
                        _artifact_page_prefix(context, index),
                        context.config.parsing.pdf_render_dpi,
                    )
                if image_path:
                    for table_number, table in enumerate(tables, start=1):
                        screenshot = _crop_table_screenshot(
                            image_path,
                            table,
                            float(plumber_page.width),
                            float(plumber_page.height),
                            image_path.parent / f"page-{index:04d}-table-{table_number:03d}.png",
                        )
                        table.screenshot_path = str(screenshot) if screenshot else None
                pages.append(
                    PageV1(
                        page_id=stable_page_id(context.document_id, index),
                        page_number=index,
                        page_type=page_type,
                        text=text,
                        markdown="\n\n".join([text, *[table.html for table in tables]]),
                        tables=tables,
                        image_path=str(image_path) if image_path else None,
                        parser_route=parser_route,
                        quality_score=score,
                        quality_signals=signals,
                        visual_required=visual_required,
                        visual_status="PENDING" if visual_required else "NOT_REQUIRED",
                    )
                )
        finally:
            plumber.close()
        return ParsedContent(
            pages=pages,
            parser_route="PDF_ROUTED",
            parser_version="pypdf",
            warnings=warnings,
            metadata={"pdf_page_count": len(reader.pages)},
        )


class DocxParser:
    name = "python-docx"
    supported_extensions = {".docx"}

    def parse(self, context: ParseContext) -> ParsedContent:
        try:
            document = WordDocument(str(context.source_path))
        except Exception as exc:
            raise ParserError(f"DOCX 打开失败：{exc}") from exc
        paragraphs = [
            paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()
        ]
        tables = []
        for index, table in enumerate(document.tables, start=1):
            rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
            tables.append(table_from_rows(f"{context.document_id}_t{index:04d}", rows))
        text_parts = paragraphs + [
            table.serialized_text for table in tables if table.serialized_text
        ]
        text = "\n".join(text_parts)
        score, signals = quality_score(text)
        page_type = classify_page(text, tables)
        complex_page = any(table.complex for table in tables)
        image_path: Path | None = None
        warnings: list[str] = []
        if complex_page:
            try:
                with tempfile.TemporaryDirectory(prefix="docflow-docx-") as temp:
                    pdf = _convert_with_libreoffice(context.source_path, Path(temp), "pdf")
                    image_path = render_pdf_page(
                        pdf,
                        1,
                        _artifact_page_prefix(context, 1),
                        context.config.parsing.pdf_render_dpi,
                    )
            except ParserError as exc:
                image_path = render_quicklook_thumbnail(
                    context.source_path,
                    _artifact_page_prefix(context, 1).with_suffix(".png"),
                )
                if image_path is None:
                    warnings.append(str(exc))
        for table in tables:
            if table.complex and image_path:
                table.screenshot_path = str(image_path)
        page = PageV1(
            page_id=stable_page_id(context.document_id, 1),
            page_number=1,
            page_type=page_type,
            text=text,
            markdown="\n\n".join(paragraphs + [table.html for table in tables]),
            tables=tables,
            image_path=str(image_path) if image_path else None,
            parser_route="NATIVE_DOCX",
            quality_score=score,
            quality_signals=signals,
            visual_required=complex_page,
            visual_status="PENDING" if complex_page else "NOT_REQUIRED",
        )
        return ParsedContent(
            pages=[page],
            parser_route="NATIVE_DOCX",
            parser_version="python-docx",
            warnings=warnings,
            metadata={"paragraph_count": len(paragraphs), "table_count": len(tables)},
        )


class SpreadsheetParser:
    name = "openpyxl"
    supported_extensions = {".xlsx"}

    def parse(self, context: ParseContext) -> ParsedContent:
        try:
            workbook = load_workbook(context.source_path, read_only=True, data_only=True)
        except Exception as exc:
            raise ParserError(f"XLSX 打开失败：{exc}") from exc
        pages: list[PageV1] = []
        skipped_empty_sheets: list[str] = []
        workbook_sheet_count = len(workbook.worksheets)
        for sheet in workbook.worksheets:
            rows: list[list[str]] = []
            for row in sheet.iter_rows():
                values = ["" if cell.value is None else str(cell.value) for cell in row]
                if any(value.strip() for value in values):
                    rows.append(values)
                if len(rows) >= 5000:
                    break
            # LibreOffice does not render empty worksheets into the converted
            # PDF. Creating visual-required pseudo pages for them causes false
            # VISUAL_IMAGE_MISSING reviews and misaligns later sheet/page pairs.
            if not rows:
                skipped_empty_sheets.append(sheet.title)
                continue
            page_number = len(pages) + 1
            table = table_from_rows(f"{context.document_id}_t{page_number:04d}", rows)
            table.complex = True
            text = f"工作表：{sheet.title}\n{table.serialized_text}"
            score, signals = quality_score(text)
            pages.append(
                PageV1(
                    page_id=stable_page_id(context.document_id, page_number),
                    page_number=page_number,
                    page_type="TABLE",
                    text=text,
                    markdown=f"## {sheet.title}\n\n{table.html}",
                    headings=[sheet.title],
                    tables=[table],
                    parser_route="NATIVE_XLSX",
                    quality_score=score,
                    quality_signals=signals,
                    visual_required=True,
                    visual_status="PENDING",
                )
            )
        workbook.close()
        if not pages:
            raise ParserError("XLSX 没有包含数据的工作表")
        warnings: list[str] = []
        try:
            with tempfile.TemporaryDirectory(prefix="docflow-xlsx-") as temp:
                pdf = _convert_with_libreoffice(context.source_path, Path(temp), "pdf")
                for page in pages:
                    image_path = render_pdf_page(
                        pdf,
                        page.page_number,
                        _artifact_page_prefix(context, page.page_number),
                        context.config.parsing.pdf_render_dpi,
                    )
                    if image_path:
                        page.image_path = str(image_path)
                        for table in page.tables:
                            table.screenshot_path = str(image_path)
        except ParserError as exc:
            preview = render_quicklook_thumbnail(
                context.source_path,
                _artifact_page_prefix(context, 1).with_suffix(".png"),
            )
            if preview:
                pages[0].image_path = str(preview)
                for table in pages[0].tables:
                    table.screenshot_path = str(preview)
            else:
                warnings.append(str(exc))
        return ParsedContent(
            pages=pages,
            parser_route="NATIVE_XLSX",
            parser_version="openpyxl",
            warnings=warnings,
            metadata={
                "sheet_count": len(pages),
                "workbook_sheet_count": workbook_sheet_count,
                "skipped_empty_sheets": skipped_empty_sheets,
            },
        )


class ImageParser:
    name = "image+ocr"
    supported_extensions = {".jpg", ".jpeg", ".png", ".tif", ".tiff", ".bmp"}

    def parse(self, context: ParseContext) -> ParsedContent:
        try:
            with Image.open(context.source_path) as image:
                width, height = image.size
        except Exception as exc:
            raise ParserError(f"图片打开失败：{exc}") from exc
        text, engine = OcrEngine().recognize(context.source_path)
        score, signals = quality_score(text)
        signals.update({"width": width, "height": height})
        page = PageV1(
            page_id=stable_page_id(context.document_id, 1),
            page_number=1,
            page_type="SCAN",
            text=text,
            markdown=text,
            image_path=str(context.source_path),
            parser_route=f"IMAGE_OCR:{engine}",
            quality_score=score,
            quality_signals=signals,
            visual_required=True,
            visual_status="PENDING",
        )
        warnings = [] if text else ["图片 OCR 无结果"]
        return ParsedContent(
            pages=[page],
            parser_route="IMAGE_OCR",
            parser_version=engine,
            warnings=warnings,
        )


class LegacyOfficeParser:
    name = "libreoffice-converter"
    supported_extensions = {".doc", ".wps", ".xls"}

    def parse(self, context: ParseContext) -> ParsedContent:
        try:
            with tempfile.TemporaryDirectory(prefix="docflow-office-") as temp:
                temp_path = Path(temp)
                if context.source_path.suffix.lower() == ".xls":
                    converted = _convert_with_libreoffice(context.source_path, temp_path, "xlsx")
                    parser = SpreadsheetParser()
                else:
                    converted = _convert_with_libreoffice(context.source_path, temp_path, "pdf")
                    parser = PdfParser()
                converted_context = ParseContext(
                    document_id=context.document_id,
                    source_file_id=context.source_file_id,
                    source_path=converted,
                    source_sha256=context.source_sha256,
                    config_version_id=context.config_version_id,
                    config=context.config,
                    artifacts=context.artifacts,
                )
                result = parser.parse(converted_context)
                result.parser_route = f"LIBREOFFICE->{result.parser_route}"
                result.parser_version = f"libreoffice+{result.parser_version}"
                result.metadata["converted_from"] = context.source_path.suffix.lower()
                return result
        except ParserError as exc:
            preview = render_quicklook_thumbnail(
                context.source_path,
                _artifact_page_prefix(context, 1).with_suffix(".png"),
            )
            if preview is None:
                raise
            text, engine = OcrEngine().recognize(preview)
            score, signals = quality_score(text)
            return ParsedContent(
                pages=[
                    PageV1(
                        page_id=stable_page_id(context.document_id, 1),
                        page_number=1,
                        page_type="TABLE"
                        if context.source_path.suffix.lower() == ".xls"
                        else "SCAN",
                        text=text,
                        markdown=text,
                        image_path=str(preview),
                        parser_route=f"QUICKLOOK_OCR:{engine}",
                        quality_score=score,
                        quality_signals=signals,
                        visual_required=True,
                        visual_status="PENDING",
                    )
                ],
                parser_route="QUICKLOOK_FALLBACK",
                parser_version=f"quicklook+{engine}",
                warnings=[f"LibreOffice 不可用，已使用 Quick Look 降级：{exc}"],
                metadata={"converted_from": context.source_path.suffix.lower()},
            )


class DoclingEnricher:
    """Optional structural enrichment. Native page boundaries remain authoritative."""

    def enrich(self, source: Path) -> tuple[str | None, str | None]:
        if importlib.util.find_spec("docling") is None:
            return None, "Docling 未安装，已使用内置结构解析"
        try:
            from docling.datamodel.accelerator_options import AcceleratorDevice, AcceleratorOptions
            from docling.datamodel.base_models import InputFormat
            from docling.datamodel.pipeline_options import PdfPipelineOptions
            from docling.document_converter import DocumentConverter, PdfFormatOption

            if source.suffix.lower() == ".pdf":
                pipeline_options = PdfPipelineOptions(
                    accelerator_options=AcceleratorOptions(device=AcceleratorDevice.CPU),
                )
                # torch.compile currently selects Inductor, which does not support MPS.
                # Docling stays on CPU; ColPali continues to use MPS independently.
                pipeline_options.layout_options.engine_options.compile_model = False
                converter = DocumentConverter(
                    format_options={
                        InputFormat.PDF: PdfFormatOption(pipeline_options=pipeline_options)
                    }
                )
            else:
                converter = DocumentConverter()
            result = converter.convert(str(source))
            return result.document.export_to_markdown(), None
        except Exception as exc:
            return None, f"Docling 增强失败：{str(exc)[:300]}"


def file_fingerprint(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as stream:
        while chunk := stream.read(1024 * 1024):
            digest.update(chunk)
    return digest.hexdigest()
