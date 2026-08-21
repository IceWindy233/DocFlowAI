from __future__ import annotations

import tempfile
from pathlib import Path

from docx import Document as WordDocument
from openpyxl import load_workbook
from PIL import Image
from pypdf import PdfReader

from docflow.core.settings import get_settings
from docflow.db.models import SourceFile
from docflow.services.parsers.common import render_pdf_page, render_quicklook_thumbnail
from docflow.services.parsers.native import _convert_with_libreoffice


class GoldenPreviewError(RuntimeError):
    pass


IMAGE_EXTENSIONS = {".jpg", ".jpeg", ".png", ".tif", ".tiff", ".bmp"}
OFFICE_EXTENSIONS = {".doc", ".docx", ".wps", ".xls", ".xlsx"}


def _source_path(source: SourceFile) -> Path:
    path = Path(source.source_path).expanduser().resolve()
    if not path.is_file():
        raise GoldenPreviewError(f"源文件不存在：{path}")
    return path


def ensure_golden_preview(source: SourceFile, page_number: int) -> Path:
    if page_number < 1:
        raise GoldenPreviewError("页码必须大于 0")
    source_path = _source_path(source)
    digest = source.sha256 or source.id
    output_prefix = (
        get_settings().artifact_root.expanduser().resolve()
        / "golden-previews"
        / digest
        / f"page-{page_number:04d}"
    )
    output = output_prefix.with_suffix(".png")
    if output.is_file():
        return output

    output.parent.mkdir(parents=True, exist_ok=True)
    suffix = source_path.suffix.lower()
    if suffix in IMAGE_EXTENSIONS:
        try:
            with Image.open(source_path) as image:
                image.seek(page_number - 1)
                image.convert("RGB").save(output, format="PNG")
        except (EOFError, OSError) as exc:
            raise GoldenPreviewError("图片页面无法渲染") from exc
        return output

    if suffix == ".pdf":
        rendered = render_pdf_page(source_path, page_number, output_prefix, dpi=144)
        if rendered:
            return rendered
        raise GoldenPreviewError("PDF 页面渲染失败，请确认已安装 Poppler")

    if suffix in OFFICE_EXTENSIONS:
        try:
            with tempfile.TemporaryDirectory(prefix="docflow-golden-preview-") as temp:
                pdf = _convert_with_libreoffice(source_path, Path(temp), "pdf")
                rendered = render_pdf_page(pdf, page_number, output_prefix, dpi=144)
        except Exception as exc:
            quicklook = (
                render_quicklook_thumbnail(source_path, output) if page_number == 1 else None
            )
            if quicklook:
                return quicklook
            raise GoldenPreviewError(f"Office 页面预览失败：{str(exc)[:300]}") from exc
        if rendered:
            return rendered
        raise GoldenPreviewError("Office 页面渲染失败")

    raise GoldenPreviewError(f"不支持预览的格式：{suffix}")


def suggested_text(source: SourceFile, page_number: int) -> str:
    """Return native text as a hint. It is never persisted as human Gold automatically."""
    source_path = _source_path(source)
    suffix = source_path.suffix.lower()
    try:
        if suffix == ".pdf":
            reader = PdfReader(str(source_path), strict=False)
            if page_number > len(reader.pages):
                return ""
            return (reader.pages[page_number - 1].extract_text() or "").strip()
        if suffix == ".docx":
            document = WordDocument(str(source_path))
            paragraphs = [item.text.strip() for item in document.paragraphs if item.text.strip()]
            tables = [
                "\n".join("\t".join(cell.text.strip() for cell in row.cells) for row in table.rows)
                for table in document.tables
            ]
            return "\n".join([*paragraphs, *tables]).strip()
        if suffix == ".xlsx":
            workbook = load_workbook(source_path, read_only=True, data_only=True)
            try:
                if page_number > len(workbook.worksheets):
                    return ""
                sheet = workbook.worksheets[page_number - 1]
                rows = []
                for row in sheet.iter_rows(values_only=True):
                    values = ["" if value is None else str(value) for value in row]
                    if any(value.strip() for value in values):
                        rows.append("\t".join(values))
                    if len(rows) >= 5000:
                        break
                return "\n".join(rows)
            finally:
                workbook.close()
    except Exception:
        return ""
    return ""
