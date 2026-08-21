from __future__ import annotations

from pathlib import Path

from docx import Document
from openpyxl import Workbook
from pypdf import PdfWriter

from docflow.domain.config import default_runtime_config
from docflow.domain.documents import NormalizedDocumentV1, PageV1
from docflow.services.parsers.base import ParseContext
from docflow.services.parsers.common import (
    infer_document_number,
    infer_document_role,
    infer_title,
    infer_version_role,
)
from docflow.services.parsers.native import (
    DoclingEnricher,
    DocxParser,
    PdfParser,
    SpreadsheetParser,
)
from docflow.services.pipeline import make_chunks
from docflow.services.storage import LocalArtifactStore


def test_docx_table_has_all_representations(tmp_path: Path) -> None:
    source = tmp_path / "复杂表格.docx"
    document = Document()
    document.add_paragraph("关于项目用地情况的请示")
    table = document.add_table(rows=9, cols=6)
    for row_index, row in enumerate(table.rows):
        for col_index, cell in enumerate(row.cells):
            cell.text = f"R{row_index}C{col_index}"
    document.save(source)
    config = default_runtime_config()
    context = ParseContext(
        document_id="doc_test",
        source_file_id="src_test",
        source_path=source,
        source_sha256="0" * 64,
        config_version_id="cfg_test",
        config=config,
        artifacts=LocalArtifactStore(tmp_path / "artifacts"),
    )
    result = DocxParser().parse(context)
    page = result.pages[0]
    assert page.tables[0].complex is True
    assert page.tables[0].html.startswith("<table>")
    assert "R1C0" in page.tables[0].serialized_text
    assert page.tables[0].cells
    assert page.visual_required is True


def test_chunks_keep_page_alignment() -> None:
    config = default_runtime_config()
    page = DocxParser  # keep import exercised by static analyzers
    normalized = NormalizedDocumentV1(
        document_id="doc_test",
        source_file_id="src_test",
        source_sha256="0" * 64,
        title="标题",
        case_id="case_test",
        parser_route="TEST",
        parser_version="1",
        config_version_id="cfg_test",
        pages=[
            {
                "page_id": "doc_test_p0001",
                "page_number": 1,
                "text": "第一段。\n第二段。",
                "markdown": "第一段。\n第二段。",
                "parser_route": "TEST",
                "quality_score": 1.0,
            }
        ],
    )
    chunks = make_chunks(normalized, config)
    assert chunks
    assert {chunk.page_id for chunk in chunks} == {"doc_test_p0001"}
    assert page is DocxParser


def test_docling_pdf_enrichment_forces_cpu_without_torch_compile(
    tmp_path: Path, monkeypatch
) -> None:
    source = tmp_path / "demo.pdf"
    source.write_bytes(b"%PDF-test")
    captured: dict[str, object] = {}

    class FakeDocument:
        def export_to_markdown(self) -> str:
            return "# parsed"

    class FakeResult:
        document = FakeDocument()

    def fake_convert(converter, path: str):
        from docling.datamodel.base_models import InputFormat

        option = converter.format_to_options[InputFormat.PDF]
        captured["device"] = option.pipeline_options.accelerator_options.device
        captured["compile_model"] = (
            option.pipeline_options.layout_options.engine_options.compile_model
        )
        captured["path"] = path
        return FakeResult()

    monkeypatch.setattr("docling.document_converter.DocumentConverter.convert", fake_convert)

    markdown, warning = DoclingEnricher().enrich(source)
    assert markdown == "# parsed"
    assert warning is None
    assert str(captured["device"]) in {"cpu", "AcceleratorDevice.CPU"}
    assert captured["compile_model"] is False
    assert captured["path"] == str(source)


def test_official_document_metadata_prefers_file_name_and_explicit_fields() -> None:
    pages = [
        {
            "page_id": "doc_test_p0001",
            "page_number": 1,
            "text": (
                "某市示例产业运营有限公司\n"
                "示例函【2027］12号\n"
                "关于示例公司参与“示范项目”建设的\n"
                "征求意见函\n"
                "根据党委会议（十三届〔2025〕31号）要求"
            ),
            "parser_route": "TEST",
            "quality_score": 1.0,
        }
    ]
    typed_pages = [PageV1.model_validate(page) for page in pages]
    file_name = "示例函〔2027〕12号、关于园区设备更新的征求意见函"
    assert infer_title(typed_pages, file_name) == "关于园区设备更新的征求意见函"
    assert (
        infer_document_number(typed_pages[0].text, f"{file_name}.pdf")
        == "示例函〔2027〕12号"
    )
    assert infer_version_role(f"{file_name}.pdf") == ("FORMAL", 0.85)


def test_official_document_metadata_uses_wrapped_form_title() -> None:
    page = PageV1(
        page_id="doc_form_p0001",
        page_number=1,
        text=(
            "某市示例公司公文稿纸\n"
            "文件编号：示例函〔2027〕12号\n"
            "文件标题：关于示例公司参与“示\n"
            "范项目”建设的征求意见函\n"
            "主送单位：有关单位"
        ),
        parser_route="TEST",
        quality_score=1.0,
    )
    assert infer_title([page], "公文稿纸-最新") == "关于示例公司参与“示范项目”建设的征求意见函"
    assert infer_document_number(page.text) == "示例函〔2027〕12号"
    assert infer_version_role("公文稿纸-最新.doc") == ("REVIEW", 0.7)


def test_spreadsheet_skips_empty_sheets_and_keeps_dense_page_numbers(
    tmp_path: Path, monkeypatch
) -> None:
    source = tmp_path / "含空白工作表.xlsx"
    workbook = Workbook()
    first = workbook.active
    first.title = "明细"
    first.append(["编号", "名称"])
    first.append(["001", "测试资产"])
    workbook.create_sheet("空白页")
    last = workbook.create_sheet("汇总")
    last.append(["项目", "金额"])
    last.append(["合计", 100])
    workbook.save(source)

    fake_pdf = tmp_path / "converted.pdf"
    fake_pdf.write_bytes(b"fake")
    monkeypatch.setattr(
        "docflow.services.parsers.native._convert_with_libreoffice",
        lambda *_: fake_pdf,
    )

    def fake_render(_pdf, page_number, output_prefix, _dpi):
        output = output_prefix.with_suffix(".png")
        output.parent.mkdir(parents=True, exist_ok=True)
        output.write_bytes(f"page-{page_number}".encode())
        return output

    monkeypatch.setattr("docflow.services.parsers.native.render_pdf_page", fake_render)
    context = ParseContext(
        document_id="doc_sheets",
        source_file_id="src_sheets",
        source_path=source,
        source_sha256="1" * 64,
        config_version_id="cfg_test",
        config=default_runtime_config(),
        artifacts=LocalArtifactStore(tmp_path / "artifacts"),
    )

    result = SpreadsheetParser().parse(context)

    assert [page.page_number for page in result.pages] == [1, 2]
    assert [page.headings for page in result.pages] == [["明细"], ["汇总"]]
    assert all(page.image_path and Path(page.image_path).exists() for page in result.pages)
    assert result.metadata == {
        "sheet_count": 2,
        "workbook_sheet_count": 3,
        "skipped_empty_sheets": ["空白页"],
    }


def test_scanned_pdf_always_requires_visual_index(tmp_path: Path, monkeypatch) -> None:
    source = tmp_path / "高质量扫描件.pdf"
    writer = PdfWriter()
    writer.add_blank_page(width=595, height=842)
    with source.open("wb") as stream:
        writer.write(stream)

    image = tmp_path / "page.png"
    image.write_bytes(b"image")
    monkeypatch.setattr(
        "docflow.services.parsers.native.render_pdf_page", lambda *_: image
    )
    monkeypatch.setattr(
        "docflow.services.parsers.native.OcrEngine.recognize",
        lambda *_: ("营业执照 法定代表人测试人员甲 " * 10, "rapidocr"),
    )
    context = ParseContext(
        document_id="doc_scan",
        source_file_id="src_scan",
        source_path=source,
        source_sha256="2" * 64,
        config_version_id="cfg_test",
        config=default_runtime_config(),
        artifacts=LocalArtifactStore(tmp_path / "artifacts"),
    )

    page = PdfParser().parse(context).pages[0]

    assert page.page_type == "SCAN"
    assert page.quality_score > 0.8
    assert page.visual_required is True
    assert page.visual_status == "PENDING"


def test_reply_role_takes_priority_over_request_and_letter_terms() -> None:
    assert (
        infer_document_role(
            "关于审议《示例市场招租方案》的请示/审计办复函.pdf",
            "关于审议《示例市场招租方案》的征求意见函",
        )
        == "REPLY"
    )
    assert (
        infer_document_role(
            "关于停车场实施收费的函/关于停车场经营管理事项的复函.pdf",
            "关于停车场经营管理事项的复函",
        )
        == "REPLY"
    )
